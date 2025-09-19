from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import List, Dict, Any, Optional, Union
from langchain.schema import HumanMessage, AIMessage, SystemMessage
import uuid
from datetime import datetime
import os
from enum import Enum
from dotenv import load_dotenv
import io
import base64
import fitz  # PyMuPDF for better PDF handling
import docx  # python-docx for Word documents
import json

# Load environment variables from .env file
load_dotenv(override=True)

# Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
from config.chatbot_config import CHATBOT_CONFIG
from utils.prompts import EXTRACTION_SYSTEM_PROMPT, EXTRACTION_USER_PROMPT

# Helper to get config value with fallback to default
def get_config_value(key: str):
    if CHATBOT_CONFIG.get(key) is not None:
        return CHATBOT_CONFIG[key]
    return CHATBOT_CONFIG.get(f"default_{key}")

extract_api = FastAPI(
    title="Stateless Extractor API", 
    description="Stateless document extraction service", 
    version="2.0.0"
)

# Enums for model providers
class ModelProvider(str, Enum):
    OPENAI = "openai"
    ANTHROPIC = "anthropic"
    GOOGLE = "google"

# Pydantic models
class DocumentContent(BaseModel):
    text: str
    images: List[str]  # Base64 encoded images
    file_type: str
    page_count: Optional[int] = None
    filename: str

class ParseDocumentResponse(BaseModel):
    content: DocumentContent
    timestamp: datetime

class ExtractRequest(BaseModel):
    instructions: str  # The user prompt
    template: str  # The template to extract into
    document_text: str  # The document text content
    document_images: Optional[List[str]] = []  # Base64 encoded images from document
    conversation_history: Optional[List[Dict[str, str]]] = []  # Previous messages
    system_prompt: Optional[str] = None
    temperature: Optional[float] = 0
    model: Optional[str] = get_config_value("model")
    provider: Optional[ModelProvider] = get_config_value("provider")
    max_tokens: Optional[int] = get_config_value("max_tokens")

class ExtractResponse(BaseModel):
    response: str
    timestamp: datetime
    tokens_used: Optional[int] = None

# LLM initialization
def get_llm(provider: "ModelProvider" = None, model: str = None, temperature: float = 0, max_tokens: int = None):
    import importlib
    provider_mapping = {
        ModelProvider.OPENAI: ("langchain_openai", "ChatOpenAI", "max_tokens", OPENAI_API_KEY),
        ModelProvider.ANTHROPIC: ("langchain_anthropic", "ChatAnthropic", "max_tokens", ANTHROPIC_API_KEY),
        ModelProvider.GOOGLE: ("langchain_google_vertexai", "ChatVertexAI", "max_output_tokens", GOOGLE_API_KEY),
    }
    if provider is None:
        provider = get_config_value("provider")
    if model is None:
        model = get_config_value("model")
    if max_tokens is None:
        max_tokens = get_config_value("max_tokens")
    if provider not in provider_mapping:
        raise ValueError(f"Unsupported provider: {provider}")
    module_name, class_name, max_tokens_param, api_key = provider_mapping[provider]
    try:
        module = importlib.import_module(module_name)
        llm_class = getattr(module, class_name)
        kwargs = {"model": model, "temperature": temperature, max_tokens_param: max_tokens, "api_key": api_key}
        return llm_class(**kwargs)
    except ImportError as e:
        raise ValueError(f"Failed to import {module_name}: {str(e)}. Make sure the package is installed.")
    except Exception as e:
        raise ValueError(f"Failed to initialize LLM for provider {provider} with model {model}: {str(e)}")

async def _process_pdf(file_bytes: io.BytesIO, filename: str) -> DocumentContent:
    """Process PDF files to extract text and images."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_content = ""
    images = []
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        # Extract text
        text_content += page.get_text() + "\n\n"
        
        # Extract images
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    images.append(f"data:image/png;base64,{img_base64}")
                pix = None
            except Exception as img_e:
                print(f"Error extracting image {img_index} from page {page_num}: {img_e}")
                continue
    
    page_count = len(doc)
    doc.close()
    
    return DocumentContent(
        text=text_content.strip(),
        images=images,
        file_type="pdf",
        page_count=page_count,
        filename=filename
    )

async def _process_word(file_bytes: io.BytesIO, filename: str) -> DocumentContent:
    """Process Word documents to extract text and images."""
    doc = docx.Document(file_bytes)
    text_content = ""
    images = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text_content += paragraph.text + "\n"
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_content += cell.text + "\t"
            text_content += "\n"
    
    # Extract images from Word document
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob
                img_base64 = base64.b64encode(image_bytes).decode('utf-8')
                content_type = image_part.content_type
                images.append(f"data:{content_type};base64,{img_base64}")
            except Exception as img_e:
                print(f"Error extracting image from Word document: {img_e}")
                continue
    
    return DocumentContent(
        text=text_content.strip(),
        images=images,
        file_type="docx",
        page_count=None,
        filename=filename
    )

async def _process_txt(file_bytes: io.BytesIO, filename: str) -> DocumentContent:
    """Process text files."""
    file_bytes.seek(0)
    text_content = file_bytes.read().decode('utf-8')
    
    return DocumentContent(
        text=text_content.strip(),
        images=[],
        file_type="txt",
        page_count=None,
        filename=filename
    )

def create_multimodal_message_content(
    instructions: str, 
    template: str, 
    document_text: str, 
    images: List[str] = None, 
    provider: ModelProvider = None
) -> Union[str, List[Dict[str, Any]]]:
    """
    Create message content that's compatible with the provider's multimodal capabilities.
    Returns either a string (for text-only) or a list of content blocks (for multimodal).
    """
    
    # Use the prompt template from prompts.py
    text_content = EXTRACTION_USER_PROMPT.format(
        instructions=instructions,
        template=template,
        document=document_text
    )
    
    # If no images or provider doesn't support multimodal, return text only
    if not images or provider not in [ModelProvider.ANTHROPIC, ModelProvider.OPENAI]:
        return text_content
    
    # For Anthropic Claude, create multimodal content blocks
    if provider == ModelProvider.ANTHROPIC:
        content_blocks = [{"type": "text", "text": text_content}]
        
        for img_base64 in images:
            # Remove the data:image/png;base64, prefix if present
            if img_base64.startswith("data:image"):
                media_type = img_base64.split(";")[0].split(":")[1]
                img_base64 = img_base64.split(",")[1]
            else:
                media_type = "image/png"
            
            content_blocks.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": img_base64
                }
            })
        
        return content_blocks
    
    # For OpenAI, create multimodal content blocks
    elif provider == ModelProvider.OPENAI:
        content_blocks = [{"type": "text", "text": text_content}]
        
        for img_base64 in images:
            content_blocks.append({
                "type": "image_url",
                "image_url": {
                    "url": img_base64 if img_base64.startswith("data:") else f"data:image/png;base64,{img_base64}"
                }
            })
        
        return content_blocks
    
    return text_content

def convert_history_to_messages(history: List[Dict[str, str]]) -> List:
    """Convert conversation history from client format to LangChain messages."""
    messages = []
    for msg in history:
        role = msg.get("role", "").lower()
        content = msg.get("content", "")
        
        if role == "user" or role == "human":
            messages.append(HumanMessage(content=content))
        elif role == "assistant" or role == "ai":
            messages.append(AIMessage(content=content))
        elif role == "system":
            messages.append(SystemMessage(content=content))
    
    return messages

# API Endpoints
@extract_api.post("/parse_document", response_model=ParseDocumentResponse)
async def parse_document(file: UploadFile = File(...)):
    """
    Parse a document and extract text and images.
    Returns the parsed content for the client to store.
    """
    try:
        # Read file content
        file_content = await file.read()
        file_bytes = io.BytesIO(file_content)
        
        # Process document based on file type
        if file.filename.lower().endswith('.pdf'):
            document_content = await _process_pdf(file_bytes, file.filename)
        elif file.filename.lower().endswith(('.docx', '.doc')):
            document_content = await _process_word(file_bytes, file.filename)
        elif file.filename.lower().endswith('.txt'):
            document_content = await _process_txt(file_bytes, file.filename)
        else:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file type. Please upload PDF, Word, or TXT files."
            )
        
        return ParseDocumentResponse(
            content=document_content,
            timestamp=datetime.now()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@extract_api.post("/extract", response_model=ExtractResponse)
async def extract(request: ExtractRequest):
    """
    Extract information from document content using the specified template.
    This is a stateless endpoint - all document data and history must be provided in the request.
    """
    try:
        # Initialize LLM
        llm = get_llm(request.provider, request.model, request.temperature, request.max_tokens)
        
        # Convert conversation history to messages
        history_messages = convert_history_to_messages(request.conversation_history)
        
        # Create multimodal content based on provider capabilities
        message_content = create_multimodal_message_content(
            instructions=request.instructions,
            template=request.template,
            document_text=request.document_text,
            images=request.document_images,
            provider=request.provider
        )
        
        # Build messages for LLM
        system_msg = request.system_prompt or EXTRACTION_SYSTEM_PROMPT
        messages = [SystemMessage(content=system_msg)]
        messages.extend(history_messages)
        
        # Add the current extraction request
        messages.append(HumanMessage(content=message_content))
        
        print(f"Processing extraction with {request.provider} - {request.model}")
        if request.document_images:
            print(f"Including {len(request.document_images)} images in multimodal request")
        
        # Invoke LLM
        llm_result = await llm.ainvoke(messages)
        response_text = llm_result.content if hasattr(llm_result, "content") else str(llm_result)
        
        # Try to get token usage if available
        tokens_used = None
        try:
            if hasattr(llm_result, 'response_metadata'):
                usage = llm_result.response_metadata.get('usage', {})
                tokens_used = usage.get('total_tokens')
        except:
            pass
        
        return ExtractResponse(
            response=response_text,
            timestamp=datetime.now(),
            tokens_used=tokens_used
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing extract request: {str(e)}")

@extract_api.get("/health")
async def health_check():
    """Check the health status of the API."""
    return {
        "status": "healthy",
        "timestamp": datetime.now(),
        "stateless": True
    }

@extract_api.get("/providers")
async def get_available_providers():
    """Get information about available LLM providers and their configurations."""
    return {
        "providers": [provider.value for provider in ModelProvider],
        "current_defaults": {
        "provider": get_config_value("provider"),
        "model": get_config_value("model"),
            "max_tokens": get_config_value("max_tokens")
        },
        "multimodal_support": {
            ModelProvider.ANTHROPIC.value: {
                "supported": True,
                "formats": ["image/png", "image/jpeg", "image/gif", "image/webp"],
                "recommended_models": ["claude-3-opus-20240229", "claude-3-sonnet-20240229", "claude-3-haiku-20240307"]
            },
            ModelProvider.OPENAI.value: {
                "supported": True,
                "formats": ["image/png", "image/jpeg", "image/gif", "image/webp"],
                "recommended_models": ["gpt-4-vision-preview", "gpt-4-turbo", "gpt-4o", "gpt-4o-mini"]
            },
            ModelProvider.GOOGLE.value: {
                "supported": False,
                "formats": [],
                "recommended_models": ["gemini-pro"]
            }
        }
    }

@extract_api.get("/")
async def root():
    """API documentation and usage examples."""
    return {
        "title": "Stateless Document Extraction API",
        "version": "2.0.0",
        "description": "Extract structured information from documents using LLMs",
        "endpoints": {
            "/parse_document": {
                "method": "POST",
                "description": "Parse a document and return text + images",
                "input": "file upload",
                "output": "DocumentContent object"
            },
            "/extract": {
                "method": "POST", 
                "description": "Extract from text and images provided in request",
                "input": "ExtractRequest with text, images, template, instructions",
                "output": "Extracted information"
            }
        },
        "usage_flow": {
            "streamlit_typical_flow": [
                "1. Client calls /parse_document with file upload",
                "2. Client stores the returned DocumentContent",
                "3. Client shows document preview to user",
                "4. User provides instructions and template",
                "5. Client calls /extract_from_parsed with stored DocumentContent",
                "6. Client stores conversation history",
                "7. For follow-up extractions, include conversation history"
            ]
        }
    }